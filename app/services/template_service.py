"""
Service de remplissage de templates Word (.docx).

L'utilisateur upload un fichier .docx contenant des balises {{CHAMP}}.
Ce service détecte automatiquement toutes les balises présentes et les
remplace avec les données de l'attribution / de la décharge.

Balises disponibles — Attestation :
  {{NOM}}              Nom de l'employé
  {{PRENOM}}           Prénom
  {{MATRICULE}}        Matricule
  {{SERVICE}}          Service / département
  {{POSTE}}            Poste / fonction
  {{DATE_JOUR}}        Date du jour (dd/mm/yyyy)
  {{DATE_ATTRIBUTION}} Date de la première attribution
  {{NB_MATERIELS}}     Nombre de matériels attribués
  {{MATERIELS_LISTE}}  Liste des matériels (une ligne par matériel)
  Pour chaque matériel (N = 1, 2, 3…) :
    {{MATERIEL_N}}     Libellé complet du matériel
    {{MARQUE_N}}       Marque
    {{MODELE_N}}       Modèle
    {{TYPE_N}}         Type (PC Portable, Écran…)
    {{SERIE_N}}        N° Série
    {{MAC_N}}          Adresse MAC
    {{ETAT_N}}         État (Neuf, Bon, Usagé…)
    {{DATE_ATTR_N}}    Date d'attribution (dd/mm/yyyy)

Balises disponibles — Décharge :
  (toutes les balises employé ci-dessus)
  {{MATERIEL}}         Libellé complet du matériel
  {{MARQUE}}           Marque
  {{MODELE}}           Modèle
  {{TYPE}}             Type
  {{SERIE}}            N° Série
  {{MAC}}              Adresse MAC
  {{ETAT_REMISE}}      État à la remise
  {{DATE_ATTRIBUTION}} Date d'attribution
  {{DATE_RESTITUTION}} Date de restitution
  {{MOTIF}}            Motif de restitution
  {{NOTES}}            Notes
"""
import re
import io
import os
from pathlib import Path
from datetime import date as dt_date, datetime

# ── Table d'aliases flexibles ─────────────────────────────────────────────────
# Clé = forme normalisée (majuscules, underscores) trouvée dans le document
# Valeur = clé canonique utilisée dans `fields`
# Permet à l'utilisateur d'écrire {{date}}, {{DATE}}, {{Date_Jour}}, {{aujourd_hui}}
# sans avoir à connaître le nom exact.
_ALIASES: dict[str, str] = {
    # Date
    "DATE":               "DATE_JOUR",
    "DATE_JOUR":          "DATE_JOUR",
    "AUJOURD_HUI":        "DATE_JOUR",
    "AUJOURD HUI":        "DATE_JOUR",
    "DATE_DU_JOUR":       "DATE_JOUR",
    # Date attribution
    "DATE_ATTRIBUTION":   "DATE_ATTRIBUTION",
    "DATE_REMISE":        "DATE_ATTRIBUTION",
    "DATE_ASSIGNATION":   "DATE_ATTRIBUTION",
    # Date récupération / restitution
    "DATE_RECUPERATION":  "DATE_RECUPERATION",
    "DATE_RESTITUTION":   "DATE_RECUPERATION",
    "DATE_RETOUR":        "DATE_RECUPERATION",
    # Employé
    "NOM":                "NOM",
    "NOM_EMPLOYE":        "NOM",
    "NOM_EMPLOYEE":       "NOM",
    "NOM_DE_FAMILLE":     "NOM",
    "PRENOM":             "PRENOM",
    "PRÉNOM":             "PRENOM",
    "PRENOM_EMPLOYE":     "PRENOM",
    "PRENOM NOM":         "PRENOM_NOM",   # prénom + nom concaténés
    "NOM_COMPLET":        "PRENOM_NOM",
    "NOM COMPLET":        "PRENOM_NOM",
    "EMPLOYE":            "PRENOM_NOM",
    "EMPLOYEE":           "PRENOM_NOM",
    "BENEFICIAIRE":       "PRENOM_NOM",
    "MATRICULE":          "MATRICULE",
    "MAT":                "MATRICULE",
    "MATRICULE_EMPLOYE":  "MATRICULE",
    "SERVICE":            "SERVICE",
    "DEPARTEMENT":        "SERVICE",
    "DÉPARTEMENT":        "SERVICE",
    "POSTE":              "POSTE",
    "FONCTION":           "POSTE",
    "TITRE":              "POSTE",
    # Matériels liste
    "MATERIELS":          "MATERIELS_LISTE",
    "MATERIELS_LISTE":    "MATERIELS_LISTE",
    "LISTE_MATERIELS":    "MATERIELS_LISTE",
    "LISTE":              "MATERIELS_LISTE",
    "EQUIPEMENTS":        "MATERIELS_LISTE",
    "EQUIPEMENT_LISTE":   "MATERIELS_LISTE",
    "NB_MATERIELS":       "NB_MATERIELS",
    "NOMBRE_MATERIELS":   "NB_MATERIELS",
    "QUANTITE":           "NB_MATERIELS",
    # Matériel singulier → slot 1
    "MATERIEL":           "MATERIEL_1",
    "EQUIPEMENT":         "MATERIEL_1",
    "MAT":                "MATERIEL_1",
    "MARQUE":             "MARQUE_1",
    "MODELE":             "MODELE_1",
    "MODÈLE":             "MODELE_1",
    "TYPE":               "TYPE_1",
    "SERIE":              "SERIE_1",
    "SÉRIE":              "SERIE_1",
    "N_SERIE":            "SERIE_1",
    "NUMERO_SERIE":       "SERIE_1",
    "NUMÉRO_SÉRIE":       "SERIE_1",
    "SN":                 "SERIE_1",
    "MAC":                "MAC_1",
    "ADRESSE_MAC":        "MAC_1",
    "MAC_ADDRESS":        "MAC_1",
    "ETAT":               "ETAT_1",
    "ÉTAT":               "ETAT_1",
    # Motif (décharge)
    "MOTIF":              "MOTIF",
    "MOTIF_RESTITUTION":  "MOTIF",
    "RAISON":             "MOTIF",
    "NOTES":              "NOTES",
    "REMARQUES":          "NOTES",
    "COMMENTAIRES":       "NOTES",
    # État remise (décharge)
    "ETAT_REMISE":        "ETAT_REMISE",
    "ÉTAT_REMISE":        "ETAT_REMISE",
    "ETAT_MATERIEL":      "ETAT_REMISE",
}

# Patterns numérotés : {{MATERIEL_2}}, {{mat_3}}, {{serie_2}}, etc.
_ALIAS_NUMBERED_PATTERNS = [
    (re.compile(r"^(?:MAT(?:ERIEL)?|EQUIPEMENT)_(\d+)$"),   "MATERIEL_{}"),
    (re.compile(r"^MARQUE_(\d+)$"),                          "MARQUE_{}"),
    (re.compile(r"^(?:MODELE|MODÈLE)_(\d+)$"),               "MODELE_{}"),
    (re.compile(r"^TYPE_(\d+)$"),                            "TYPE_{}"),
    (re.compile(r"^(?:SERIE|SÉRIE|N_SERIE|SN)_(\d+)$"),      "SERIE_{}"),
    (re.compile(r"^(?:MAC|ADRESSE_MAC)_(\d+)$"),             "MAC_{}"),
    (re.compile(r"^(?:ETAT|ÉTAT)_(\d+)$"),                   "ETAT_{}"),
    (re.compile(r"^DATE_ATTR_(\d+)$"),                       "DATE_ATTR_{}"),
]


def _normalize_key(raw: str) -> str:
    """Normalise une clé brute de template vers la clé canonique.

    Exemple : 'date' → 'DATE_JOUR', 'Matériel_2' → 'MATERIEL_2', 'NOM' → 'NOM'
    """
    key = raw.strip().upper().replace(" ", "_").replace("-", "_")
    # Alias exact
    if key in _ALIASES:
        return _ALIASES[key]
    # Alias numéroté
    for pattern, template_fmt in _ALIAS_NUMBERED_PATTERNS:
        m = pattern.match(key)
        if m:
            return template_fmt.format(m.group(1))
    # Clé déjà canonique
    return key

TEMPLATE_DIR = Path(__file__).parent.parent.parent / "templates"
TEMPLATE_DIR.mkdir(exist_ok=True)

# Chemin des templates stockés
TEMPLATE_PATHS = {
    "attestation":  TEMPLATE_DIR / "attestation.docx",
    "decharge":     TEMPLATE_DIR / "decharge.docx",
    "recuperation": TEMPLATE_DIR / "recuperation.docx",
}

TYPE_LABELS = {
    "ORDINATEUR_PORTABLE": "PC Portable",
    "ORDINATEUR_FIXE":     "PC Bureau",
    "ECRAN":               "Écran",
    "SOURIS":              "Souris",
    "CLAVIER":             "Clavier",
    "TELEPHONE":           "Téléphone",
    "TABLETTE":            "Tablette",
    "IMPRIMANTE":          "Imprimante",
    "SWITCH":              "Switch réseau",
    "ROUTEUR":             "Routeur",
    "ONDULEUR":            "Onduleur",
    "AP":                  "Point d'accès",
    "SERVEUR":             "Serveur",
    "PARE_FEU":            "Pare-feu",
    "AUTRE":               "Matériel informatique",
}


# ── Utilitaires ───────────────────────────────────────────────────────────────

def _fmt_date(d) -> str:
    if not d:
        return ""
    if isinstance(d, (dt_date, datetime)):
        return d.strftime("%d/%m/%Y")
    return str(d)


def _mat_label(m) -> str:
    """Libellé court d'un matériel : 'PC Portable Dell Latitude 5420'."""
    parts = [TYPE_LABELS.get(m.type_materiel, m.type_materiel), m.marque]
    if m.modele:
        parts.append(m.modele)
    return " ".join(parts)


def scan_placeholders(docx_bytes: bytes) -> list[dict]:
    """Extrait toutes les balises {{...}} présentes dans le template.

    Retourne une liste de dicts {raw, canonical} pour afficher la correspondance
    automatique à l'utilisateur (ex: raw='date', canonical='DATE_JOUR').
    """
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    # Pattern flexible : accepte lettres, chiffres, espaces, tirets, underscores, accents
    pattern = re.compile(r"\{\{([^}]{1,60})\}\}")
    found: dict[str, str] = {}  # raw → canonical

    def scan_para(para):
        full = "".join(r.text for r in para.runs)
        for m in pattern.finditer(full):
            raw = m.group(1).strip()
            canonical = _normalize_key(raw)
            found[raw] = canonical

    for para in doc.paragraphs:
        scan_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    scan_para(para)
    for section in doc.sections:
        for para in section.header.paragraphs:
            scan_para(para)
        for para in section.footer.paragraphs:
            scan_para(para)

    return [{"raw": k, "canonical": v} for k, v in sorted(found.items())]


def _replace_in_doc(doc, fields: dict[str, str]):
    """Remplace toutes les balises {{...}} dans le document par leur valeur.

    La clé brute est normalisée via _normalize_key avant de chercher dans fields.
    Les paragraphes dont TOUT le contenu est une balise dont la valeur est vide
    sont supprimés.
    """
    # Pattern flexible : tout ce qui est entre {{ et }}
    pattern      = re.compile(r"\{\{([^}]{1,60})\}\}")
    only_pattern = re.compile(r"^\s*\{\{([^}]{1,60})\}\}\s*$")

    def _resolve(raw_key: str) -> str | None:
        """Retourne la valeur pour une clé brute, ou None si inconnue."""
        return fields.get(_normalize_key(raw_key))

    def _replace_para(para) -> bool:
        """Retourne True si le paragraphe doit être supprimé (balise → valeur vide)."""
        if not para.runs:
            return False
        full_text = "".join(r.text for r in para.runs)
        m = only_pattern.match(full_text)
        if m:
            val = _resolve(m.group(1))
            if val is not None and val.strip() == "":
                return True
        new_text = pattern.sub(
            lambda mx: (v if (v := _resolve(mx.group(1))) is not None else mx.group(0)),
            full_text,
        )
        if new_text != full_text:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""
        return False

    def _remove_para(para):
        """Supprime physiquement un paragraphe du XML."""
        p = para._p
        p.getparent().remove(p)

    to_remove = []
    for para in doc.paragraphs:
        if _replace_para(para):
            to_remove.append(para)
    for para in to_remove:
        _remove_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_remove = []
                for para in cell.paragraphs:
                    if _replace_para(para):
                        cell_remove.append(para)
                for para in cell_remove:
                    _remove_para(para)

    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_para(para)
        for para in section.footer.paragraphs:
            _replace_para(para)


# ── Générateurs ───────────────────────────────────────────────────────────────

def generate_attestation_from_template(attributions: list) -> bytes:
    """
    Génère l'attestation en remplissant le template uploadé.
    `attributions` : liste d'attributions ACTIVES d'un même employé.
    """
    from docx import Document
    template_path = TEMPLATE_PATHS["attestation"]
    if not template_path.exists():
        raise FileNotFoundError("Aucun template d'attestation uploadé")

    doc = Document(str(template_path))
    a0  = attributions[0]  # référence employé

    prenom_nom = f"{a0.employee_prenom or ''} {a0.employee_nom or ''}".strip()
    fields: dict[str, str] = {
        "NOM":              a0.employee_nom or "",
        "PRENOM":           a0.employee_prenom or "",
        "PRENOM_NOM":       prenom_nom,
        "MATRICULE":        a0.employee_matricule or "",
        "SERVICE":          a0.employee_service or "",
        "POSTE":            a0.employee_poste or "",
        "DATE_JOUR":        datetime.now().strftime("%d/%m/%Y"),
        "DATE_ATTRIBUTION": _fmt_date(min(a.date_attribution for a in attributions)),
        "NB_MATERIELS":     str(len(attributions)),
        "MATERIELS_LISTE":  "\n".join(
            f"• {_mat_label(a.materiel)}" for a in attributions if a.materiel
        ),
    }

    # Remplir les slots de matériels (1 à 10) — les slots vides seront supprimés
    for i in range(1, 11):
        if i <= len(attributions):
            a_i = attributions[i - 1]
            m   = a_i.materiel
            if m:
                # Construire la description complète selon le type d'identifiant
                desc = _mat_label(m)
                if m.adresse_mac:
                    desc += f", Mac Address : {m.adresse_mac}"
                elif m.numero_serie:
                    desc += f", S/N : {m.numero_serie}"
                fields[f"MATERIEL_{i}"]  = desc
                fields[f"MARQUE_{i}"]    = m.marque or ""
                fields[f"MODELE_{i}"]    = m.modele or ""
                fields[f"TYPE_{i}"]      = TYPE_LABELS.get(m.type_materiel, m.type_materiel or "")
                fields[f"SERIE_{i}"]     = m.numero_serie or ""
                fields[f"MAC_{i}"]       = m.adresse_mac or ""
                fields[f"ETAT_{i}"]      = m.etat or ""
                fields[f"DATE_ATTR_{i}"] = _fmt_date(a_i.date_attribution)
            else:
                fields[f"MATERIEL_{i}"] = ""
        else:
            # Slot vide → sera supprimé par _replace_in_doc
            fields[f"MATERIEL_{i}"] = ""

    _replace_in_doc(doc, fields)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_decharge_from_template(attribution) -> bytes:
    """
    Génère la décharge en remplissant le template uploadé.
    `attribution` : une attribution (active ou clôturée).
    """
    from docx import Document
    template_path = TEMPLATE_PATHS["decharge"]
    if not template_path.exists():
        raise FileNotFoundError("Aucun template de décharge uploadé")

    doc = Document(str(template_path))
    a = attribution
    m = a.materiel

    MOTIF_LABELS = {
        "DEPART": "Départ", "CHANGEMENT": "Changement", "PANNE": "Panne",
        "FIN_CONTRAT": "Fin de contrat", "AUTRE": "Autre",
    }

    prenom_nom_d = f"{a.employee_prenom or ''} {a.employee_nom or ''}".strip()
    fields: dict[str, str] = {
        "NOM":               a.employee_nom or "",
        "PRENOM":            a.employee_prenom or "",
        "PRENOM_NOM":        prenom_nom_d,
        "MATRICULE":         a.employee_matricule or "",
        "SERVICE":           a.employee_service or "",
        "POSTE":             a.employee_poste or "",
        "DATE_JOUR":         datetime.now().strftime("%d/%m/%Y"),
        "DATE_ATTRIBUTION":  _fmt_date(a.date_attribution),
        "DATE_RECUPERATION": _fmt_date(a.date_restitution),
        "DATE_RESTITUTION":  _fmt_date(a.date_restitution),
        "MOTIF":             MOTIF_LABELS.get(a.motif_restitution or "", a.motif_restitution or ""),
        "ETAT_REMISE":       a.etat_remise or "",
        "NOTES":             a.notes or "",
        # Matériel (singulier → slot 1)
        "MATERIEL_1": _mat_label(m) if m else "",
        "MARQUE_1":   m.marque or "" if m else "",
        "MODELE_1":   m.modele or "" if m else "",
        "TYPE_1":     TYPE_LABELS.get(m.type_materiel, m.type_materiel or "") if m else "",
        "SERIE_1":    m.numero_serie or "" if m else "",
        "MAC_1":      m.adresse_mac or "" if m else "",
    }

    _replace_in_doc(doc, fields)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_recuperation_from_template(attributions: list, date_recuperation=None) -> bytes:
    """
    Génère l'attestation de récupération en remplissant le template uploadé.

    Balises disponibles :
      {{NOM}}, {{PRENOM}}, {{MATRICULE}}, {{SERVICE}}, {{POSTE}}
      {{DATE_JOUR}}, {{DATE_RECUPERATION}}, {{NB_MATERIELS}}, {{MATERIELS_LISTE}}
      {{MATERIEL_N}}, {{MARQUE_N}}, {{MODELE_N}}, {{TYPE_N}}, {{SERIE_N}}, {{MAC_N}}
    """
    from docx import Document
    template_path = TEMPLATE_PATHS["recuperation"]
    if not template_path.exists():
        raise FileNotFoundError("Aucun template de récupération uploadé")

    doc = Document(str(template_path))
    a0  = attributions[0]

    date_recup = date_recuperation or datetime.now().date()
    if hasattr(date_recup, "strftime"):
        date_recup_str = date_recup.strftime("%d/%m/%Y")
    else:
        date_recup_str = str(date_recup)

    prenom_nom_r = f"{a0.employee_prenom or ''} {a0.employee_nom or ''}".strip()
    fields: dict[str, str] = {
        "NOM":               a0.employee_nom or "",
        "PRENOM":            a0.employee_prenom or "",
        "PRENOM_NOM":        prenom_nom_r,
        "MATRICULE":         a0.employee_matricule or "",
        "SERVICE":           a0.employee_service or "",
        "POSTE":             a0.employee_poste or "",
        "DATE_JOUR":         datetime.now().strftime("%d/%m/%Y"),
        "DATE_RECUPERATION": date_recup_str,
        "NB_MATERIELS":      str(len(attributions)),
        "MATERIELS_LISTE":   "\n".join(
            f"• {_mat_label(a.materiel)}" for a in attributions if a.materiel
        ),
    }

    for i in range(1, 11):
        if i <= len(attributions):
            a_i = attributions[i - 1]
            m   = a_i.materiel
            if m:
                desc = _mat_label(m)
                if m.adresse_mac:
                    desc += f", Mac Address : {m.adresse_mac}"
                elif m.numero_serie:
                    desc += f", S/N : {m.numero_serie}"
                fields[f"MATERIEL_{i}"] = desc
                fields[f"MARQUE_{i}"]   = m.marque or ""
                fields[f"MODELE_{i}"]   = m.modele or ""
                fields[f"TYPE_{i}"]     = TYPE_LABELS.get(m.type_materiel, m.type_materiel or "")
                fields[f"SERIE_{i}"]    = m.numero_serie or ""
                fields[f"MAC_{i}"]      = m.adresse_mac or ""
            else:
                fields[f"MATERIEL_{i}"] = ""
        else:
            fields[f"MATERIEL_{i}"] = ""

    _replace_in_doc(doc, fields)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def template_exists(doc_type: str) -> bool:
    return TEMPLATE_PATHS.get(doc_type, Path("")).exists()


def save_template(doc_type: str, content: bytes) -> None:
    path = TEMPLATE_PATHS.get(doc_type)
    if not path:
        raise ValueError(f"Type inconnu : {doc_type}")
    path.write_bytes(content)
