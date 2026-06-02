"""
Génère le template Word d'attestation de mise à disposition CAMUSAT SÉNÉGAL.
Exécuter avec : python generate_attestation_template.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = Path("/app/templates/attestation.docx")
OUT.parent.mkdir(exist_ok=True)

BLUE  = RGBColor(0x00, 0x3C, 0x71)  # bleu CAMUSAT
RED   = RGBColor(0xC0, 0x39, 0x2B)
BLACK = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# ── Marges ────────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        val = kwargs.get(side)
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   val.get('val','none'))
            border.set(qn('w:sz'),    str(val.get('sz', 0)))
            border.set(qn('w:space'),'0')
            border.set(qn('w:color'), val.get('color','auto'))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def no_border_table(tbl):
    """Retire toutes les bordures d'un tableau."""
    tbl.style = 'Table Grid'
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={'val':'none'}, bottom={'val':'none'},
                left={'val':'none'}, right={'val':'none'})


def add_run(para, text, bold=False, italic=False, underline=False,
            size=10.5, color=BLACK):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def add_para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        add_run(p, text)
    return p


def add_hline(doc, color="003C71", width=12):
    """Ajoute une ligne horizontale via une bordure de paragraphe."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(width))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# 1. EN-TÊTE
# ══════════════════════════════════════════════════════════════════════════════
hdr = doc.add_table(rows=1, cols=3)
hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_border_table(hdr)
hdr.columns[0].width = Cm(5)
hdr.columns[1].width = Cm(9)
hdr.columns[2].width = Cm(5)

# Colonne gauche — nom société
c0 = hdr.cell(0, 0)
p0 = c0.paragraphs[0]
r0 = add_run(p0, "CAMUSAT SENEGAL", bold=True, size=9, color=BLUE)
c0.add_paragraph()  # espace
p0b = c0.add_paragraph()
add_run(p0b, "00 (221) 33 865 10 01", size=7, color=BLUE)
p0c = c0.add_paragraph()
add_run(p0c, "contact.senegal@camusat.com", size=7, color=BLUE)
p0d = c0.add_paragraph()
add_run(p0d, "www.camusat.com", size=7, color=BLUE)

# Colonne centrale — adresse
c1 = hdr.cell(0, 1)
p1 = c1.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p1, "10083 Sacré Cœur 3, VDN - Immeuble\n"
            "Sounin 3 - 1er Étage, Dakar BP 45923,\n"
            "Sénégal",
        size=7, color=BLUE)

# Colonne droite — (logo texte)
c2 = hdr.cell(0, 2)
p2 = c2.paragraphs[0]
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Ligne rouge sous l'en-tête
add_hline(doc, color="C0392B", width=16)

# ══════════════════════════════════════════════════════════════════════════════
# 2. DATE
# ══════════════════════════════════════════════════════════════════════════════
p_date = add_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=8, space_after=4)
add_run(p_date, "Date, le ", size=10.5)
add_run(p_date, "{{DATE_JOUR}}", bold=True, size=10.5)

# ══════════════════════════════════════════════════════════════════════════════
# 3. OBJET
# ══════════════════════════════════════════════════════════════════════════════
p_obj = add_para(doc, space_before=8, space_after=8)
add_run(p_obj, "Objet : ", bold=True, underline=True, size=10.5)
add_run(p_obj, "Attestation de mise à disposition de matériels informatiques",
        underline=True, size=10.5)

# ══════════════════════════════════════════════════════════════════════════════
# 4. CORPS — intro
# ══════════════════════════════════════════════════════════════════════════════
p_intro = add_para(doc, space_before=4, space_after=6)
add_run(p_intro, "Je soussigné(e) ", size=10.5)
add_run(p_intro, "{{PRENOM}} {{NOM}}", bold=True, size=10.5)
add_run(p_intro, " certifie avoir reçu les matériels suivants :", size=10.5)

# ── Matériels (bullet slots 1 à 8) ───────────────────────────────────────────
for i in range(1, 9):
    p_mat = doc.add_paragraph(style='List Bullet')
    p_mat.paragraph_format.space_before = Pt(1)
    p_mat.paragraph_format.space_after  = Pt(1)
    p_mat.paragraph_format.left_indent  = Cm(0.8)
    add_run(p_mat, f"{{{{MATERIEL_{i}}}}}", size=10.5)

# ══════════════════════════════════════════════════════════════════════════════
# 5. PARAGRAPHE RAPPEL
# ══════════════════════════════════════════════════════════════════════════════
p_rappel = add_para(doc, space_before=10, space_after=6)
p_rappel.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p_rappel,
    "Nous vous rappelons que les matériels suivants susvisés mise à votre disposition ce jour "
    "ne peut en aucun cas être utilisé pour des raisons personnelles. Son usage doit en effet "
    "rester strictement professionnel.",
    size=10.5)

# ══════════════════════════════════════════════════════════════════════════════
# 6. PARAGRAPHE RESTITUTION
# ══════════════════════════════════════════════════════════════════════════════
p_rest = add_para(doc, space_before=6, space_after=6)
p_rest.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p_rest,
    "Par la présente attestation contre signée de remise du matériel informatique, vous vous "
    "engagez à restituer ses matériels à CAMUSAT SENEGAL à la fin de votre contrat de travail.",
    size=10.5)

# ══════════════════════════════════════════════════════════════════════════════
# 7. FORMULE DE POLITESSE
# ══════════════════════════════════════════════════════════════════════════════
p_pol = add_para(doc, space_before=6, space_after=14)
p_pol.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p_pol,
    "Nous vous prions d'agréer, Monsieur/Madame, nos salutations distinguées.",
    size=10.5)

# ══════════════════════════════════════════════════════════════════════════════
# 8. SIGNATURES
# ══════════════════════════════════════════════════════════════════════════════
sig_tbl = doc.add_table(rows=3, cols=2)
no_border_table(sig_tbl)
sig_tbl.columns[0].width = Cm(9)
sig_tbl.columns[1].width = Cm(9)

# Ligne 1 — titres
p_sig_l = sig_tbl.cell(0, 0).paragraphs[0]
add_run(p_sig_l, "L'intéressé(e) :", bold=True, underline=True, size=10.5)

p_sig_r = sig_tbl.cell(0, 1).paragraphs[0]
p_sig_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p_sig_r, "L'Assistant Informatique", bold=True, underline=True, size=10.5)

# Ligne 2 — espace signature
sig_tbl.cell(1, 0).paragraphs[0].paragraph_format.space_after = Pt(24)
p_r2 = sig_tbl.cell(1, 1).paragraphs[0]
p_r2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_r2.paragraph_format.space_after = Pt(24)

# Ligne 3 — nom signataire droit
p_r3 = sig_tbl.cell(2, 1).paragraphs[0]
p_r3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p_r3, "Libasse SARR", bold=True, size=10.5)

# ══════════════════════════════════════════════════════════════════════════════
# 9. PIED DE PAGE — ligne rouge + logo texte
# ══════════════════════════════════════════════════════════════════════════════
add_hline(doc, color="C0392B", width=16)

p_foot = add_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=4)
add_run(p_foot, "camusat", bold=True, size=14, color=BLUE)

# ══════════════════════════════════════════════════════════════════════════════
doc.save(str(OUT))
print(f"✓ Template généré : {OUT}")
